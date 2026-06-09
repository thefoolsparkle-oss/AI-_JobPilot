import os
from pathlib import Path

from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent.parent


class DocumentExportService:
    def render_resume(self, resume_data: dict, template_path: str) -> str:
        full_template_path = self._resolve_template(template_path)
        doc = Document(full_template_path)

        def replace_placeholders(text: str) -> str:
            for key, value in resume_data.items():
                placeholder = "{{" + key + "}}"
                if placeholder in text:
                    value_str = str(value) if not isinstance(value, (list, dict)) else ""
                    text = text.replace(placeholder, value_str)
            return text

        for paragraph in doc.paragraphs:
            full_text = paragraph.text
            replaced = replace_placeholders(full_text)
            if replaced != full_text and paragraph.runs:
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = replaced

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = paragraph.text
                        replaced = replace_placeholders(full_text)
                        if replaced != full_text and paragraph.runs:
                            for run in paragraph.runs:
                                run.text = ""
                            paragraph.runs[0].text = replaced

        output_dir = os.path.join(BASE_DIR, "output")
        os.makedirs(output_dir, exist_ok=True)
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(output_dir, f"resume_{timestamp}.docx")
        doc.save(output_path)
        return output_path

    def render_resume_html(self, resume_data: dict) -> str:
        name = resume_data.get("name", "姓名")
        phone = resume_data.get("phone", "")
        email = resume_data.get("email", "")
        location = resume_data.get("location", "")
        github = resume_data.get("github", "")
        linkedin = resume_data.get("linkedin", "")

        section_defs = [
            ("求职意向", resume_data.get("objective", "")),
            ("专业优势", resume_data.get("summary", "")),
            ("实习经历", resume_data.get("internships", "")),
            ("项目经历", resume_data.get("projects", "")),
            ("教育背景", resume_data.get("education", "")),
            ("技能", resume_data.get("skills", "")),
        ]

        sections_html = ""
        for title, content in section_defs:
            if not content:
                continue
            lines = [ln.strip().lstrip("- ") for ln in str(content).strip().split("\n") if ln.strip()]
            if not lines:
                continue
            lis = "".join(f"<li>{line}</li>" for line in lines)
            sections_html += f'<div class="section"><div class="section-title">{title}</div><ul>{lis}</ul></div>\n'

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm 2.5cm; }}
  body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; font-size: 10pt; color: #1a1a1a; line-height: 1.6; }}
  .header {{ text-align: center; margin-bottom: 20px; }}
  .header h1 {{ font-size: 18pt; margin: 0 0 4px 0; }}
  .header .contact {{ font-size: 9pt; color: #555; }}
  .section {{ margin-bottom: 16px; }}
  .section-title {{ font-size: 12pt; font-weight: bold; color: #1d4ed8; border-bottom: 1px solid #bfdbfe; padding-bottom: 4px; margin-bottom: 8px; }}
  ul {{ margin: 0; padding-left: 18px; }}
  li {{ margin-bottom: 4px; }}
</style>
</head>
<body>
<div class="header">
  <h1>{name}</h1>
  <div class="contact">{phone} | {email} | {location}</div>
  <div class="contact">{github} | {linkedin}</div>
</div>
{sections_html}
</body>
</html>"""

        output_dir = os.path.join(BASE_DIR, "output")
        os.makedirs(output_dir, exist_ok=True)
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        html_path = os.path.join(output_dir, f"resume_{timestamp}.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)
        return html_path

    async def export_pdf(self, html_path: str) -> str:
        try:
            from playwright.async_api import async_playwright
        except ImportError:
            raise RuntimeError("Playwright not installed. Run: pip install playwright && playwright install chromium")

        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.goto(f"file:///{html_path.replace(chr(92), '/')}")
            await page.wait_for_load_state("networkidle")
            pdf_path = html_path.replace(".html", ".pdf")
            await page.pdf(path=pdf_path, format="A4", margin={"top": "2cm", "bottom": "2cm", "left": "2.5cm", "right": "2.5cm"})
            await browser.close()
            return pdf_path

    def _resolve_template(self, template_path: str) -> str:
        if os.path.isfile(template_path):
            return template_path
        full = os.path.join(BASE_DIR, template_path)
        if os.path.isfile(full):
            return full
        full = os.path.join(BASE_DIR, "..", template_path)
        if os.path.isfile(full):
            return os.path.abspath(full)
        return template_path
